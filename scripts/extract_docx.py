import sys
from docx import Document

def docx_to_text(infile):
    doc = Document(infile)
    texts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    texts.append(cell.text.strip())
    return "\n".join(texts)

def main():
    if len(sys.argv) >= 3:
        infile = sys.argv[1]
        outfile = sys.argv[2]
    else:
        infile = r"D:\dev\新手外贸销售实战手册.docx"
        outfile = r"D:\dev\新手外贸销售实战手册.txt"
    try:
        text = docx_to_text(infile)
        with open(outfile, 'w', encoding='utf-8') as f:
            f.write(text)
        print('WROTE ' + outfile)
    except Exception as e:
        print('ERROR:', e)

if __name__ == '__main__':
    main()
